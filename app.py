import argparse
import os
import sqlite3
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

from fastapi import FastAPI, HTTPException, Response, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from graph import build_graph

# Export tools
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document


app = FastAPI(title="MSRA API", version="2.2")

# ---- CORS (frontend can call API) ----
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # for demo; lock down later
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

graph = build_graph()

DB_PATH = os.getenv("MSRA_DB_PATH", "msra_chats.db")

MAX_CHARS_PER_MESSAGE = 6000
DEFAULT_CHAT_TITLE = "New Chat"

# -----------------------
# Frontend paths
# -----------------------
BASE_DIR = Path(__file__).resolve().parent
FRONTEND_DIR = BASE_DIR / "frontend"
STATIC_DIR = FRONTEND_DIR / "static"

# Serve static files at /static
if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


# -----------------------
# DB helpers
# -----------------------
def _db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn
'''
def _db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL;")   # ✅ ADD THIS
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn

'''


def _init_db() -> None:
    conn = _db()
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS chats (
        id TEXT PRIMARY KEY,
        title TEXT NOT NULL,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS messages (
        id TEXT PRIMARY KEY,
        chat_id TEXT NOT NULL,
        role TEXT NOT NULL,            -- 'user' or 'assistant'
        content TEXT NOT NULL,
        refs_json TEXT,                -- JSON list[str] for assistant messages
        created_at TEXT NOT NULL,
        FOREIGN KEY(chat_id) REFERENCES chats(id) ON DELETE CASCADE
    )
    """)

    cur.execute("CREATE INDEX IF NOT EXISTS idx_messages_chat ON messages(chat_id)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_messages_chat_time ON messages(chat_id, created_at)")
    conn.commit()
    conn.close()


_init_db()


def _now() -> str:
    return datetime.utcnow().isoformat()


def _json_dumps_refs(refs: List[str]) -> str:
    import json
    return json.dumps(refs or [], ensure_ascii=False)


def _json_loads_refs(s: Optional[str]) -> List[str]:
    import json
    if not s:
        return []
    try:
        out = json.loads(s)
        return out if isinstance(out, list) else []
    except Exception:
        return []


def _unique_preserve_order(items: List[str]) -> List[str]:
    seen = set()
    out: List[str] = []
    for x in items or []:
        if isinstance(x, str):
            xx = x.strip()
            if xx and xx not in seen:
                seen.add(xx)
                out.append(xx)
    return out


# -----------------------
# Models
# -----------------------
class ChatMeta(BaseModel):
    id: str
    title: str
    updated_at: str


class Message(BaseModel):
    id: str
    role: str
    content: str
    references: List[str] = []
    created_at: str


class CreateChatRequest(BaseModel):
    title: Optional[str] = None


class RenameChatRequest(BaseModel):
    title: str


class SendMessageRequest(BaseModel):
    message: str


class SendMessageResponse(BaseModel):
    chat_id: str
    assistant_message: Message
    all_messages: List[Message]


# -----------------------
# Frontend routes
# -----------------------
@app.get("/", response_class=HTMLResponse)
def home_page():
    home_file = FRONTEND_DIR / "home.html"
    if home_file.exists():
        return FileResponse(str(home_file))
    return HTMLResponse("<h2>Home not found</h2><p>Create frontend/home.html</p>", status_code=200)


@app.get("/ui", response_class=HTMLResponse)
def ui_page():
    ui_file = FRONTEND_DIR / "ui.html"
    if ui_file.exists():
        return FileResponse(str(ui_file))
    return HTMLResponse("<h2>UI not found</h2><p>Create frontend/ui.html</p>", status_code=200)


# -----------------------
# Pipeline runner
# -----------------------
def _run_pipeline(query: str) -> Tuple[str, List[str]]:
    q = (query or "").strip()
    if not q:
        return "Query is empty.", []

    if len(q) > MAX_CHARS_PER_MESSAGE:
        q = q[:MAX_CHARS_PER_MESSAGE]

    final_state: Dict[str, Any] = graph.invoke({"user_query": q})

    answer = (final_state.get("final_answer") or "").strip()
    refs = _unique_preserve_order(final_state.get("references") or [])

    if not answer:
        answer = "No answer produced."

    return answer, refs


# -----------------------
# CRUD helpers
# -----------------------
def _chat_exists(conn: sqlite3.Connection, chat_id: str) -> bool:
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM chats WHERE id = ?", (chat_id,))
    return cur.fetchone() is not None


def _fetch_messages(conn: sqlite3.Connection, chat_id: str) -> List[Message]:
    cur = conn.cursor()
    cur.execute("""
      SELECT id, role, content, refs_json, created_at
      FROM messages
      WHERE chat_id = ?
      ORDER BY created_at ASC
    """, (chat_id,))
    rows = cur.fetchall()
    out: List[Message] = []
    for r in rows:
        out.append(Message(
            id=r["id"],
            role=r["role"],
            content=r["content"],
            references=_json_loads_refs(r["refs_json"]),
            created_at=r["created_at"],
        ))
    return out


def _fetch_message_row(conn: sqlite3.Connection, chat_id: str, message_id: str) -> Optional[sqlite3.Row]:
    cur = conn.cursor()
    cur.execute("""
      SELECT id, role, content, refs_json, created_at
      FROM messages
      WHERE chat_id = ? AND id = ?
      LIMIT 1
    """, (chat_id, message_id))
    return cur.fetchone()


def _fetch_next_assistant_after(conn: sqlite3.Connection, chat_id: str, created_at: str) -> Optional[sqlite3.Row]:
    cur = conn.cursor()
    cur.execute("""
      SELECT id, role, content, refs_json, created_at
      FROM messages
      WHERE chat_id = ?
        AND role = 'assistant'
        AND created_at > ?
      ORDER BY created_at ASC
      LIMIT 1
    """, (chat_id, created_at))
    return cur.fetchone()


# -----------------------
# API
# -----------------------
@app.get("/health")
def health() -> Dict[str, str]:
    return {"status": "ok"}


@app.get("/chats", response_model=List[ChatMeta])
def list_chats() -> List[ChatMeta]:
    conn = _db()
    cur = conn.cursor()
    cur.execute("""
      SELECT id, title, updated_at
      FROM chats
      ORDER BY updated_at DESC
    """)
    rows = cur.fetchall()
    conn.close()
    return [ChatMeta(id=r["id"], title=r["title"], updated_at=r["updated_at"]) for r in rows]


@app.post("/chats", response_model=ChatMeta)
def create_chat(req: CreateChatRequest) -> ChatMeta:
    chat_id = uuid.uuid4().hex
    title = (req.title or "").strip() or DEFAULT_CHAT_TITLE
    ts = _now()

    conn = _db()
    conn.execute(
        "INSERT INTO chats (id, title, created_at, updated_at) VALUES (?, ?, ?, ?)",
        (chat_id, title, ts, ts)
    )
    conn.commit()
    conn.close()
    return ChatMeta(id=chat_id, title=title, updated_at=ts)


@app.get("/chats/{chat_id}", response_model=List[Message])
def get_chat_messages(chat_id: str) -> List[Message]:
    conn = _db()
    if not _chat_exists(conn, chat_id):
        conn.close()
        raise HTTPException(status_code=404, detail="Chat not found")
    msgs = _fetch_messages(conn, chat_id)
    conn.close()
    return msgs


@app.patch("/chats/{chat_id}", response_model=ChatMeta)
def rename_chat(chat_id: str, req: RenameChatRequest) -> ChatMeta:
    title = (req.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title cannot be empty")

    conn = _db()
    if not _chat_exists(conn, chat_id):
        conn.close()
        raise HTTPException(status_code=404, detail="Chat not found")

    ts = _now()
    conn.execute("UPDATE chats SET title = ?, updated_at = ? WHERE id = ?", (title, ts, chat_id))
    conn.commit()
    conn.close()
    return ChatMeta(id=chat_id, title=title, updated_at=ts)


@app.delete("/chats/{chat_id}")
def delete_chat(chat_id: str) -> Dict[str, str]:
    conn = _db()
    conn.execute("DELETE FROM chats WHERE id = ?", (chat_id,))
    conn.commit()
    conn.close()
    return {"status": "ok"}


@app.post("/chats/{chat_id}/messages", response_model=SendMessageResponse)
def send_message(chat_id: str, req: SendMessageRequest) -> SendMessageResponse:
    user_msg = (req.message or "").strip()
    if not user_msg:
        raise HTTPException(status_code=400, detail="Message is empty")

    conn = _db()
    if not _chat_exists(conn, chat_id):
        conn.close()
        raise HTTPException(status_code=404, detail="Chat not found")

    # store user message
    user_id = uuid.uuid4().hex
    ts_user = _now()
    conn.execute(
        "INSERT INTO messages (id, chat_id, role, content, refs_json, created_at) VALUES (?, ?, ?, ?, ?, ?)",
        (user_id, chat_id, "user", user_msg, None, ts_user)
    )

    # run MSRA
    answer, refs = _run_pipeline(user_msg)

    # store assistant message
    asst_id = uuid.uuid4().hex
    ts_asst = _now()
    conn.execute(
        "INSERT INTO messages (id, chat_id, role, content, refs_json, created_at) VALUES (?, ?, ?, ?, ?, ?)",
        (asst_id, chat_id, "assistant", answer, _json_dumps_refs(refs), ts_asst)
    )

    # bump chat updated_at
    conn.execute("UPDATE chats SET updated_at = ? WHERE id = ?", (ts_asst, chat_id))

    conn.commit()

    all_msgs = _fetch_messages(conn, chat_id)
    conn.close()

    assistant_message = Message(
        id=asst_id,
        role="assistant",
        content=answer,
        references=refs,
        created_at=ts_asst
    )

    return SendMessageResponse(chat_id=chat_id, assistant_message=assistant_message, all_messages=all_msgs)


# -----------------------
# Export helpers
# -----------------------
def _build_export_text(messages: List[Message]) -> str:
    lines: List[str] = []
    for m in messages:
        if m.role == "user":
            lines.append(f"User:\n{m.content}")
        else:
            lines.append(f"Assistant:\n{m.content}")
            if m.references:
                lines.append("\nReferences:")
                for r in m.references:
                    lines.append(f"- {r}")
        lines.append("")
    return "\n".join(lines).strip()


def _export_bytes_pdf(title: str, export_text: str) -> bytes:
    from io import BytesIO
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    x = 40
    y = height - 50
    c.setTitle(title)

    def draw_wrapped(text: str):
        nonlocal y
        max_width = width - 80
        for paragraph in text.split("\n"):
            words = paragraph.split(" ")
            line = ""
            for w in words:
                test = (line + " " + w).strip()
                if c.stringWidth(test, "Helvetica", 10) <= max_width:
                    line = test
                else:
                    c.setFont("Helvetica", 10)
                    c.drawString(x, y, line)
                    y -= 14
                    line = w
                    if y < 60:
                        c.showPage()
                        y = height - 50
            c.setFont("Helvetica", 10)
            c.drawString(x, y, line)
            y -= 14
            if y < 60:
                c.showPage()
                y = height - 50
        y -= 8

    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, title)
    y -= 24

    draw_wrapped(export_text)
    c.save()
    pdf_bytes = buf.getvalue()
    buf.close()
    return pdf_bytes


def _export_bytes_docx(title: str, messages: List[Message]) -> bytes:
    from io import BytesIO
    doc = Document()
    doc.add_heading(title, level=1)

    for m in messages:
        if m.role == "user":
            doc.add_paragraph("User:", style="Heading 3")
            doc.add_paragraph(m.content)
        else:
            doc.add_paragraph("Assistant:", style="Heading 3")
            doc.add_paragraph(m.content)
            if m.references:
                doc.add_paragraph("References:", style="Heading 4")
                for r in m.references:
                    doc.add_paragraph(r, style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    buf.close()
    return data


def _safe_filename(title: str) -> str:
    safe = "".join(ch for ch in (title or "") if ch.isalnum() or ch in (" ", "_", "-")).strip() or "export"
    return safe.replace(" ", "_")


# -----------------------
# Export (whole chat) - keep as-is
# -----------------------
@app.get("/chats/{chat_id}/export")
def export_chat(
    chat_id: str,
    format: str = Query("pdf", pattern="^(pdf|docx)$")
):
    conn = _db()
    if not _chat_exists(conn, chat_id):
        conn.close()
        raise HTTPException(status_code=404, detail="Chat not found")

    row = conn.execute("SELECT title FROM chats WHERE id = ?", (chat_id,)).fetchone()
    title = row["title"] if row else "chat"

    messages = _fetch_messages(conn, chat_id)
    conn.close()

    export_text = _build_export_text(messages)
    safe_title = _safe_filename(title)

    if format == "pdf":
        pdf_bytes = _export_bytes_pdf(title, export_text)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'}
        )

    docx_bytes = _export_bytes_docx(title, messages)
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'}
    )


# -----------------------
# Export (single turn): one user query + its next assistant response ✅
# -----------------------
@app.get("/chats/{chat_id}/turns/{user_message_id}/export")
def export_turn(
    chat_id: str,
    user_message_id: str,
    format: str = Query("pdf", pattern="^(pdf|docx)$"),
):
    conn = _db()
    if not _chat_exists(conn, chat_id):
        conn.close()
        raise HTTPException(status_code=404, detail="Chat not found")

    user_row = _fetch_message_row(conn, chat_id, user_message_id)
    if not user_row or user_row["role"] != "user":
        conn.close()
        raise HTTPException(status_code=404, detail="User message not found")

    asst_row = _fetch_next_assistant_after(conn, chat_id, user_row["created_at"])
    conn.close()

    # Build the turn as messages
    user_msg = Message(
        id=user_row["id"],
        role=user_row["role"],
        content=user_row["content"],
        references=[],
        created_at=user_row["created_at"],
    )

    if asst_row:
        asst_msg = Message(
            id=asst_row["id"],
            role=asst_row["role"],
            content=asst_row["content"],
            references=_json_loads_refs(asst_row["refs_json"]),
            created_at=asst_row["created_at"],
        )
        turn_msgs = [user_msg, asst_msg]
    else:
        # If somehow assistant not present, export only user query
        turn_msgs = [user_msg]

    # Title for export file
    short_q = (user_msg.content or "").strip().replace("\n", " ")
    short_q = short_q[:80] + ("..." if len(short_q) > 80 else "")
    title = f"MSRA - {short_q}" if short_q else "MSRA - Export"
    safe_title = _safe_filename(title)

    export_text = _build_export_text(turn_msgs)

    if format == "pdf":
        pdf_bytes = _export_bytes_pdf(title, export_text)
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.pdf"'}
        )

    docx_bytes = _export_bytes_docx(title, turn_msgs)
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'}
    )


# -----------------------
# CLI loop (optional)
# -----------------------
def main_cli():
    parser = argparse.ArgumentParser(description="MSRA server (FastAPI) + quick CLI")
    parser.add_argument("--serve", action="store_true", help="Hint to run uvicorn.")
    parser.add_argument("--query", type=str, help="Run one query (no DB chat).")
    args = parser.parse_args()

    if args.serve:
        print("Run:\n  uvicorn app:app --reload --workers 1\n")
        return

    if args.query:
        ans, refs = _run_pipeline(args.query)
        print("\n========== FINAL OUTPUT ==========\n")
        print(ans)
        if refs:
            print("\nReferences:")
            for r in refs:
                print("-", r)
        print("\n========== END ==========\n")
        return

    print("For UI, run the API server with uvicorn.")
    print("  uvicorn app:app --reload --workers 1")


if __name__ == "__main__":
    main_cli()
#uvicorn app:app --reload --workers 1